import io
from pathlib import Path

import pytest
from docx import Document

from config import settings
from services.extract_text_from_document import extract_text_from_document
from utils.errors import DocumentExtractError, FileSizeError


def _minimal_docx_bytes() -> bytes:
    buf = io.BytesIO()
    doc = Document()
    doc.add_paragraph("x" * max(settings.MIN_EXTRACTED_TEXT_CHARS + 10, 50))
    doc.save(buf)
    return buf.getvalue()


def test_extract_text_from_document_docx_roundtrip(tmp_path: Path):
    path = tmp_path / "r.docx"
    path.write_bytes(_minimal_docx_bytes())
    text = extract_text_from_document(path)
    assert len(text.strip()) >= settings.MIN_EXTRACTED_TEXT_CHARS


def test_extract_text_from_document_docx_empty_raises(tmp_path: Path):
    buf = io.BytesIO()
    doc = Document()
    doc.save(buf)
    path = tmp_path / "empty.docx"
    path.write_bytes(buf.getvalue())
    with pytest.raises(DocumentExtractError):
        extract_text_from_document(path)


def test_extract_text_from_document_rejects_tiny_file_on_disk(tmp_path: Path):
    path = tmp_path / "tiny.pdf"
    path.write_bytes(b"%PDF-1.4\n")
    with pytest.raises(FileSizeError):
        extract_text_from_document(path)
