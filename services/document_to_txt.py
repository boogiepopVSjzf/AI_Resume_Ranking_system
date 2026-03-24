from __future__ import annotations

from pathlib import Path

from docx import Document

from services.document_validate import validate_file_size
from services.pdf_to_txt import pdf_to_txt
from services.text_clean_service import finalize_extracted_plaintext
from utils.errors import DocumentExtractError
from utils.logger import get_logger

logger = get_logger("document_extract")


def document_to_txt(path: Path) -> str:
    validate_file_size(path)
    ext = path.suffix.lower()
    if ext == ".pdf":
        return pdf_to_txt(path, skip_size_check=True)
    if ext == ".docx":
        return _docx_to_txt(path)
    raise DocumentExtractError(f"不支持的文件类型: {ext}")


def _docx_to_txt(path: Path) -> str:
    try:
        doc = Document(str(path))
    except Exception as exc:
        logger.error("DOCX open failed: %s — %s", path.name, exc)
        raise DocumentExtractError(
            f"无法打开或解析 DOCX：{path.name}（文件可能已损坏或已加密）"
        ) from exc

    parts: list[str] = []
    for paragraph in doc.paragraphs:
        t = (paragraph.text or "").strip()
        if t:
            parts.append(t)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                t = (cell.text or "").strip()
                if t:
                    parts.append(t)

    raw = "\n".join(parts)
    return finalize_extracted_plaintext(raw, source="docx")
